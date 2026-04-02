import os
import sys
import subprocess
import argparse

# 自动安装依赖
def install_dependencies():
    required_packages = ["python-docx", "pandas", "openpyxl", "python-dotenv", "volcengine-python-sdk", "httpx", "pydantic"]
    for package in required_packages:
        try:
            __import__(package.replace("-", "_"))
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "--break-system-packages", package])

install_dependencies()

from docx import Document
import pandas as pd
import json
from volcenginesdkarkruntime import Ark
from dotenv import load_dotenv

load_dotenv()

# 初始化豆包大模型客户端
client = Ark(
    api_key=os.getenv("VOLC_API_KEY"),
    base_url=os.getenv("VOLC_BASE_URL", "https://ark.cn-beijing.volces.com/api/v3")
)
MODEL = os.getenv("VOLC_MODEL", "doubao-seed-2.0-pro")

# 读取UAT规范
with open(os.path.join(os.path.dirname(__file__), "../references/bank_uat_standard.md"), "r", encoding="utf-8") as f:
    UAT_STANDARD = f.read()

def read_docx(file_path):
    """读取Word文档内容"""
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append(" | ".join(row_text))
    return "\n".join(full_text)

def generate_test_cases(requirement_content, file_name):
    """调用大模型生成测试案例"""
    prompt = f"""
你是专业的银行业UAT测试案例编写专家，请根据以下业务需求文档内容，严格按照《银行业UAT测试案例编写规范》生成测试案例。

# 需求文档来源：{file_name}
# 需求内容：
{requirement_content}

# 编写规范：
{UAT_STANDARD}

请输出JSON格式的测试案例列表，每个案例包含以下字段：
- scenario: 测试场景
- case_id: 测试案例ID（格式：UAT-XXX-001，XXX为模块缩写，数字自动递增）
- precondition: 前置条件
- steps: 测试步骤（数组格式，每个步骤为字符串）
- expected: 预期结果
- priority: 优先级（高/中/低）
- requirement_source: 需求来源（就是当前的文件名）
- remark: 备注（可选）

只输出JSON，不要其他任何说明文字。
"""
    response = client.chat.completions.create(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

def main():
    global MODEL
    parser = argparse.ArgumentParser(description="业务需求文档转UAT测试案例工具")
    parser.add_argument("input_path", help="Word文件路径或者包含Word文件的目录路径")
    parser.add_argument("output_path", help="输出Excel文件路径")
    parser.add_argument("--template", help="自定义Excel模板路径", default=None)
    parser.add_argument("--model", help="使用的大模型ID", default=MODEL)
    args = parser.parse_args()

    MODEL = args.model

    # 收集所有Word文件
    docx_files = []
    if os.path.isdir(args.input_path):
        for root, dirs, files in os.walk(args.input_path):
            for file in files:
                if file.endswith(".docx") and not file.startswith("~$"):
                    docx_files.append(os.path.join(root, file))
    else:
        if args.input_path.endswith(".docx"):
            docx_files.append(args.input_path)
        else:
            print("错误：输入路径不是Word文件或目录")
            sys.exit(1)

    if not docx_files:
        print("错误：未找到任何Word文档")
        sys.exit(1)

    all_cases = []
    for file in docx_files:
        print(f"处理文件：{file}")
        content = read_docx(file)
        cases = generate_test_cases(content, os.path.basename(file))
        if isinstance(cases, list):
            all_cases.extend(cases)
        else:
            all_cases.append(cases)

    # 转换为DataFrame
    df = pd.DataFrame(all_cases)
    # 调整列顺序
    columns = ["scenario", "case_id", "precondition", "steps", "expected", "priority", "requirement_source", "remark"]
    df = df[columns]

    # 写入Excel
    if args.template and os.path.exists(args.template):
        with pd.ExcelWriter(args.output_path, engine="openpyxl", mode="a", if_sheet_exists="replace") as writer:
            df.to_excel(writer, sheet_name="测试案例", index=False)
    else:
        df.to_excel(args.output_path, sheet_name="测试案例", index=False)

    print(f"生成完成！共生成{len(all_cases)}条测试案例，输出文件：{args.output_path}")

if __name__ == "__main__":
    main()